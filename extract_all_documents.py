#!/usr/bin/env python3
"""
Extract text from all downloaded call documents and cache as JSON.

Processes all ./downloads/{call_code}/ directories:
- Unzips ZIP files
- Extracts text from PDFs (via pymupdf)
- Extracts text from DOCXs (via python-docx)
- Saves to cache/documents/{call_code}.json

Usage:
    python3 extract_all_documents.py
"""

import json
import os
import re
import sys
import zipfile
import tempfile
import shutil
from pathlib import Path

import fitz  # pymupdf
from docx import Document

DOWNLOADS_DIR = Path("downloads")
CACHE_DIR = Path("cache/documents")
EXTRACTED_DIR = Path("extracted")


def extract_zip(zip_path: Path, dest_dir: Path) -> list:
    """Extract a ZIP file, handling Slovak encoding. Returns list of extracted file paths."""
    extracted = []
    try:
        with zipfile.ZipFile(zip_path, 'r') as z:
            for info in z.infolist():
                # Fix encoding
                try:
                    name = info.filename.encode('cp437').decode('utf-8')
                except (UnicodeDecodeError, UnicodeEncodeError):
                    name = info.filename

                basename = os.path.basename(name)
                if not basename or basename.startswith('~$'):
                    continue

                # Sanitize filename
                basename = re.sub(r'[<>:"/\\|?*]', '_', basename)
                target = dest_dir / basename

                with z.open(info) as f_in:
                    with open(target, 'wb') as f_out:
                        f_out.write(f_in.read())
                extracted.append(target)
    except (zipfile.BadZipFile, Exception) as e:
        print(f"    ⚠️  ZIP error {zip_path.name}: {e}")
    return extracted


def extract_text_from_pdf(path: Path) -> str:
    """Extract text from a PDF file."""
    try:
        doc = fitz.open(str(path))
        pages = [page.get_text() for page in doc]
        doc.close()
        text = "\n\n".join(pages)
        return re.sub(r'\n{3,}', '\n\n', text).strip()
    except Exception as e:
        print(f"    ⚠️  PDF error {path.name}: {e}")
        return ""


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip()
    except Exception as e:
        print(f"    ⚠️  DOCX error {path.name}: {e}")
        return ""


def process_call(call_code: str, call_dir: Path) -> dict:
    """Process all documents for a single call. Returns dict of doc_name → {filename, text, chars}."""
    docs = {}

    # First, collect all files (direct + from ZIPs)
    temp_dir = Path(tempfile.mkdtemp())
    all_files = []

    for f in call_dir.iterdir():
        if f.suffix.lower() == '.zip':
            # Extract ZIP contents to temp dir
            extracted = extract_zip(f, temp_dir)
            all_files.extend(extracted)
        elif f.suffix.lower() in ('.pdf', '.docx', '.doc'):
            all_files.append(f)

    # Now extract text from each file
    for fpath in all_files:
        text = ""
        fname = fpath.name

        if fname.lower().endswith('.pdf'):
            text = extract_text_from_pdf(fpath)
        elif fname.lower().endswith('.docx'):
            text = extract_text_from_docx(fpath)
        elif fname.lower().endswith('.doc'):
            continue  # Can't read .doc without antiword/libreoffice
        else:
            continue

        if text and len(text) > 50:  # Skip near-empty extractions
            doc_name = fname.rsplit('.', 1)[0]
            # Avoid duplicates (keep longer version)
            if doc_name in docs and len(text) <= docs[doc_name]['chars']:
                continue
            docs[doc_name] = {
                "filename": fname,
                "call_code": call_code,
                "text": text,
                "chars": len(text),
            }

    # Cleanup temp dir
    shutil.rmtree(temp_dir, ignore_errors=True)

    return docs


def main():
    CACHE_DIR.mkdir(parents=True, exist_ok=True)

    # Find all downloaded call directories
    call_dirs = sorted([d for d in DOWNLOADS_DIR.iterdir() if d.is_dir()])
    print(f"Found {len(call_dirs)} call directories to process\n")

    total_docs = 0
    total_chars = 0
    processed = 0
    skipped = 0

    for i, call_dir in enumerate(call_dirs):
        call_code = call_dir.name
        cache_file = CACHE_DIR / f"{call_code}.json"

        # Skip if already cached
        if cache_file.exists():
            with open(cache_file) as f:
                existing = json.load(f)
            if existing:
                skipped += 1
                total_docs += len(existing)
                total_chars += sum(d['chars'] for d in existing.values())
                print(f"  [{i+1}/{len(call_dirs)}] {call_code}: cached ({len(existing)} docs)")
                continue

        print(f"  [{i+1}/{len(call_dirs)}] {call_code}: extracting...", end=" ")

        docs = process_call(call_code, call_dir)

        if docs:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(docs, f, ensure_ascii=False, indent=1)

            chars = sum(d['chars'] for d in docs.values())
            total_docs += len(docs)
            total_chars += chars
            processed += 1
            print(f"{len(docs)} docs, {chars:,} chars")
        else:
            print("no extractable text")

    print(f"\n{'='*60}")
    print(f"DONE: {processed} calls processed, {skipped} already cached")
    print(f"Total: {total_docs} documents, {total_chars:,} characters")
    print(f"Cache: {sum(1 for _ in CACHE_DIR.glob('*.json'))} files in {CACHE_DIR}/")


if __name__ == "__main__":
    main()
